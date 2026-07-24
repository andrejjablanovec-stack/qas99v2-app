import streamlit as st
from openai import OpenAI
import json
import pandas as pd
from datetime import datetime


# =====================================================
# NASTAVITVE STRANI
# =====================================================

st.set_page_config(
    page_title="QAS-99 Evalvacija",
    page_icon="📋",
    layout="wide"
)


# =====================================================
# CSS
# =====================================================

st.markdown(
"""
<style>

.main {
    padding-top: 1rem;
}

h1 {
    color: #1f4e79;
}

.stButton > button {
    background-color: #1f4e79;
    color:white;
    font-weight:600;
    border-radius:8px;
    border:none;
}

.stButton > button:hover {
    background-color:#2c6aa0;
    color:white;
}


.info-box {
    background:#eef5fb;
    border-left:6px solid #1f4e79;
    padding:20px;
    border-radius:8px;
}


.warning-box {
    background:#fff7ef;
    border-left:6px solid #fd7e14;
    padding:20px;
    border-radius:8px;
}


.success-box {
    background:#f1fff4;
    border-left:6px solid #198754;
    padding:20px;
    border-radius:8px;
}


.metric-box {
    padding:15px;
    border-radius:8px;
    background:#f7f7f7;
}


table {
    width:100%;
    border-collapse:collapse;
}


th {
    background:#1f4e79;
    color:white;
    padding:10px;
}


td {
    border:1px solid #ddd;
    padding:8px;
}


</style>
""",
unsafe_allow_html=True
)



# =====================================================
# OPENAI / GROQ
# =====================================================

client = OpenAI(
    api_key=st.secrets["GROQ_API_KEY"],
    base_url="https://api.groq.com/openai/v1"
)



# =====================================================
# QAS POSTAVKE
# =====================================================

QAS_ITEMS = {

"Q1a":"Določanje, kateri deli vprašanja se preberejo",
"Q1b":"Manjkajoče informacije",
"Q1c":"Zahtevnost branja",

"Q2a":"Napačna ali nasprotujoča si navodila",
"Q2b":"Zapletena navodila",

"Q3a":"Dolgo ali nerodno besedilo",
"Q3b":"Nejasni strokovni izrazi",
"Q3c":"Dvoumno vprašanje",
"Q3d":"Nejasno referenčno obdobje",

"Q4a":"Neustrezne predpostavke o respondentu",
"Q4b":"Predpostavka stalnega vedenja",
"Q4c":"Dvojno vprašanje (double-barreled)",

"Q5a":"Pomanjkanje znanja",
"Q5b":"Neoblikovano mnenje",
"Q5c":"Težaven priklic iz spomina",
"Q5d":"Zahtevno miselno računanje",

"Q6a":"Občutljiva tema",
"Q6b":"Neprimerno besedilo za občutljivo temo",
"Q6c":"Družbeno zaželen odgovor",

"Q7a":"Neprimerno odprto vprašanje",
"Q7b":"Neujemanje vprašanja in odgovornih kategorij",
"Q7c":"Nejasne odgovorne kategorije",
"Q7d":"Dvoumne odgovorne kategorije",
"Q7e":"Prekrivajoče se odgovorne kategorije",
"Q7f":"Manjkajoče odgovorne kategorije",
"Q7g":"Nelogičen vrstni red odgovorov",

"Q8a":"Druge metodološke težave"

}



# =====================================================
# SYSTEM PROMPT
# =====================================================

SYSTEM_PROMPT = f"""

Ste strokovnjak za metodologijo anketiranja.
Ocenjujete vprašanja po RTI Question Appraisal System (QAS-99).

Vaša naloga je simulirati metodološki pregled vprašalnika.

Pri ocenjevanju upoštevajte:

- besedilo vprašanja,
- odgovorne kategorije,
- ciljno populacijo,
- način zbiranja podatkov,
- različne vrste respondentov.


Pri vsakem kriteriju določite:

problem:
DA ali NE

resnost:
- Visoka
- Srednja
- Nizka
- Brez težave


Visoka:
Težava lahko povzroči sistematično napačne podatke.

Srednja:
Del respondentov lahko vprašanje napačno razume.

Nizka:
Manjša izboljšava formulacije.


Če je problem NE:
razlaga mora biti "/".

Če je problem DA:
podajte kratko metodološko razlago.


Uporabite naslednje QAS kriterije:

{json.dumps(QAS_ITEMS, ensure_ascii=False, indent=2)}



====================================================
DODATNA METODOLOŠKA PRAVILA
====================================================

Pri presoji se vživite v:

- starejšega respondenta,
- respondenta z nižjo izobrazbo,
- respondenta brez strokovnega znanja,
- respondenta, ki vprašanje bere sam.


Bodite previdni.
Če obstaja realna možnost napačne interpretacije,
raje označite DA.


====================================================
IZHOD
====================================================

Odgovorite IZKLJUČNO v veljavnem JSON formatu.


Struktura:

{{
"summary": {{
    "quality_score": 1-10,
    "confidence": "Visoka/Srednja/Nizka",
    "main_problems":[
        "problem 1",
        "problem 2"
    ]
}},

"items":[

{{
"code":"Q3c",
"name":"Dvoumno vprašanje",
"problem":"DA",
"severity":"Srednja",
"explanation":"..."
}}

],

"additional_comments":
"...",

"suggestion":
"..."

}}


Ne dodajajte nobenega drugega besedila izven JSON.
"""



# =====================================================
# SESSION STATE
# =====================================================

if "history" not in st.session_state:
    st.session_state.history = []

# =====================================================
# NASLOV
# =====================================================

st.markdown(
"""
<h1 style="text-align:center;">
📋 QAS-99 Evalvacija vprašanj
</h1>
""",
unsafe_allow_html=True
)


st.markdown(
"""
<div class="info-box">

<b>RTI Question Appraisal System (QAS-99)</b>

<br><br>

Orodje uporablja veliki jezikovni model za metodološki pregled
anketnih vprašanj. Analiza preverja razumljivost, kognitivno
obremenitev, predpostavke, občutljivost teme in kakovost odgovornih
kategorij.

</div>
""",
unsafe_allow_html=True
)



# =====================================================
# TABELE
# =====================================================

tab1, tab2 = st.tabs(
[
"📋 Evalvacija vprašanja",
"ℹ️ O aplikaciji"
]
)



# =====================================================
# GLAVNI TAB
# =====================================================

with tab1:


    st.subheader("Vnos vprašanja")


    col1, col2 = st.columns(2)


    with col1:

        question = st.text_area(
            "Besedilo vprašanja",
            height=150,
            placeholder=
            """
Primer:

Kako pogosto uporabljate spletne storitve javne uprave?
"""
        )


    with col2:

        categories = st.text_area(
            "Odgovorne kategorije",
            height=150,
            placeholder=
            """
Vsak dan
Večkrat tedensko
Enkrat tedensko
Redkeje
Nikoli
"""
        )



    st.divider()



    st.subheader("Kontekst raziskave")


    col1, col2 = st.columns(2)


    with col1:

        population = st.text_input(
            "Ciljna populacija (opcijsko)",
            placeholder=
            "npr. prebivalci Slovenije, zaposleni, pacienti"
        )


        survey_mode = st.selectbox(
            "Način zbiranja podatkov",
            [
                "Spletna anketa (CAWI)",
                "Telefonska anketa (CATI)",
                "Osebno anketiranje (CAPI)",
                "Papirni vprašalnik (PAPI)"
            ]
        )


    with col2:

        research_goal = st.text_area(
            "Namen raziskave (opcijsko)",
            height=120,
            placeholder=
            """
Kaj želite z vprašanjem izmeriti?
"""
        )



    st.divider()



    if st.button(
        "🔍 Analiziraj vprašanje",
        use_container_width=True
    ):


        if not question.strip():

            st.warning(
                "Najprej vnesite vprašanje."
            )

            st.stop()



        user_prompt = f"""

ANALIZA VPRAŠANJA


VPRAŠANJE:

{question}



ODGOVORNE KATEGORIJE:

{
categories 
if categories.strip()
else
"Odprto vprašanje"
}



KONTEKST:


Ciljna populacija:

{population if population else "Ni podano"}



Način zbiranja:

{survey_mode}



Namen raziskave:

{research_goal if research_goal else "Ni podan"}

"""


        with st.spinner(
            "🔍 Izvajam metodološko evalvacijo..."
        ):


            try:


                response = client.chat.completions.create(

                    model="llama-3.3-70b-versatile",

                    temperature=0.1,

                    messages=[

                        {
                        "role":"system",
                        "content":SYSTEM_PROMPT
                        },

                        {
                        "role":"user",
                        "content":user_prompt
                        }

                    ]

                )


                raw_result = response.choices[0].message.content



                # odstranimo morebitne markdown oznake

                raw_result = (
                    raw_result
                    .replace("```json","")
                    .replace("```","")
                    .strip()
                )



                result = json.loads(raw_result)



                # shranimo zgodovino

                st.session_state.history.append(
                    {
                    "čas":
                    datetime.now().strftime(
                        "%d.%m.%Y %H:%M"
                    ),

                    "vprašanje":
                    question,

                    "rezultat":
                    result
                    }
                )



                st.success(
                    "Analiza uspešno zaključena."
                )



                # =====================================
                # POVZETEK
                # =====================================


                st.subheader(
                    "📊 Povzetek evalvacije"
                )


                summary = result["summary"]



                c1,c2,c3 = st.columns(3)



                with c1:

                    st.metric(
                        "Metodološka kakovost",
                        f"{summary['quality_score']}/10"
                    )


                with c2:

                    st.metric(
                        "Zaupanje modela",
                        summary["confidence"]
                    )


                with c3:

                    st.metric(
                        "Število zaznanih težav",
                        len(
                            [
                            x for x in result["items"]
                            if x["problem"]=="DA"
                            ]
                        )
                    )



                if summary["main_problems"]:


                    st.markdown(
                    """
                    <div class="warning-box">

                    <b>Glavne ugotovljene težave:</b>

                    </div>
                    """,
                    unsafe_allow_html=True
                    )


                    for problem in summary["main_problems"]:

                        st.write(
                            "⚠️ " + problem
                        )




                st.divider()



                # =====================================
                # QAS TABELA
                # =====================================


                st.subheader(
                    "📋 Podrobna QAS analiza"
                )


                table_data=[]


                for item in result["items"]:


                    table_data.append(

                    {

                    "Postavka":
                    item["code"],


                    "Ime":
                    item["name"],


                    "Težava":
                    item["problem"],


                    "Resnost":
                    item["severity"],


                    "Razlaga":
                    item["explanation"]

                    }

                    )



                df = pd.DataFrame(table_data)



                def highlight_problem(row):

                    if row["Težava"]=="DA":

                        return [
                        "background-color:#ffe6e6"
                        ]*len(row)

                    else:

                        return [
                        "background-color:#e9f7ef"
                        ]*len(row)



                st.dataframe(

                    df.style.apply(
                        highlight_problem,
                        axis=1
                    ),

                    use_container_width=True,

                    height=700

                )



                st.divider()



                # =====================================
                # KOMENTARJI
                # =====================================


                st.subheader(
                    "📝 Dodatni komentarji"
                )


                st.info(
                    result["additional_comments"]
                )



                st.subheader(
                    "💡 Predlog izboljšanega vprašanja"
                )


                st.success(
                    result["suggestion"]
                )



            except Exception as e:


                st.error(
                    f"Napaka pri analizi: {e}"
                )

# =====================================================
# ZGODOVINA ANALIZ
# =====================================================

with tab2:


    st.subheader(
        "ℹ️ O aplikaciji"
    )


    st.markdown(
    """
    <div class="info-box">

    <b>QAS-99 Evaluator</b>

    <br><br>

    Aplikacija uporablja metodo
    <b>Question Appraisal System (QAS-99)</b>
    za sistematično preverjanje kakovosti anketnih vprašanj.

    <br><br>

    Preverja:

    <ul>
    <li>jasnost vprašanja</li>
    <li>kognitivno zahtevnost</li>
    <li>predpostavke</li>
    <li>težave pri priklicu informacij</li>
    <li>občutljivost teme</li>
    <li>ustreznost odgovornih kategorij</li>
    </ul>

    </div>
    """,
    unsafe_allow_html=True
    )


    st.divider()


    st.subheader(
        "🕒 Zgodovina analiz"
    )


    if not st.session_state.history:


        st.info(
            "V tej seji še ni bilo izvedenih analiz."
        )


    else:


        for i, analysis in enumerate(
            reversed(st.session_state.history)
        ):


            with st.expander(
                f"{analysis['čas']} - {analysis['vprašanje'][:80]}"
            ):


                st.write(
                    analysis["vprašanje"]
                )


                result = analysis["rezultat"]


                st.metric(
                    "Ocena kakovosti",
                    f"{result['summary']['quality_score']}/10"
                )


                problems = [

                    x for x in result["items"]

                    if x["problem"]=="DA"

                ]


                st.write(
                    f"Zaznanih težav: {len(problems)}"
                )



# =====================================================
# FUNKCIJE ZA IZVOZ
# =====================================================


def create_excel(result):


    rows=[]


    for item in result["items"]:


        rows.append(

            {

            "Postavka":
            item["code"],

            "Ime":
            item["name"],

            "Težava":
            item["problem"],

            "Resnost":
            item["severity"],

            "Razlaga":
            item["explanation"]

            }

        )


    df=pd.DataFrame(rows)


    return df



def create_word(result, question):


    from docx import Document
    from io import BytesIO


    doc=Document()


    doc.add_heading(
        "QAS-99 Evalvacija vprašanja",
        level=1
    )


    doc.add_heading(
        "Vprašanje",
        level=2
    )


    doc.add_paragraph(
        question
    )


    doc.add_heading(
        "Povzetek",
        level=2
    )


    doc.add_paragraph(

        f"""
Ocena kakovosti:
{result['summary']['quality_score']}/10

Zaupanje:
{result['summary']['confidence']}
"""

    )


    doc.add_heading(
        "Rezultati QAS",
        level=2
    )


    table = doc.add_table(
        rows=1,
        cols=5
    )


    headers=[
        "Postavka",
        "Ime",
        "Težava",
        "Resnost",
        "Razlaga"
    ]


    for i,h in enumerate(headers):

        table.rows[0].cells[i].text=h



    for item in result["items"]:


        row=table.add_row().cells


        row[0].text=item["code"]

        row[1].text=item["name"]

        row[2].text=item["problem"]

        row[3].text=item["severity"]

        row[4].text=item["explanation"]



    buffer=BytesIO()


    doc.save(buffer)


    buffer.seek(0)


    return buffer



def create_pdf(result, question):


    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer
    )

    from reportlab.lib.styles import getSampleStyleSheet

    from io import BytesIO


    buffer=BytesIO()


    doc=SimpleDocTemplate(
        buffer
    )


    styles=getSampleStyleSheet()


    content=[]


    content.append(

        Paragraph(
            "QAS-99 Evalvacija vprašanja",
            styles["Title"]
        )

    )


    content.append(
        Spacer(1,20)
    )


    content.append(

        Paragraph(
            "Vprašanje:",
            styles["Heading2"]
        )

    )


    content.append(

        Paragraph(
            question,
            styles["BodyText"]
        )

    )


    content.append(
        Spacer(1,20)
    )


    content.append(

        Paragraph(

            f"""
Ocena kakovosti:
{result['summary']['quality_score']}/10
<br/>
Zaupanje:
{result['summary']['confidence']}
""",

            styles["BodyText"]

        )

    )


    doc.build(
        content
    )


    buffer.seek(0)


    return buffer



# =====================================================
# IZVOZ ZADNJE ANALIZE
# =====================================================

if st.session_state.history:


    latest = st.session_state.history[-1]


    result = latest["rezultat"]


    question = latest["vprašanje"]


    st.sidebar.header(
        "📥 Izvoz rezultatov"
    )


    excel_file=create_excel(
        result
    )


    st.sidebar.download_button(

        label="📊 Prenesi Excel",

        data=
        excel_file.to_csv(
            index=False
        ),

        file_name=
        "QAS99_evalvacija.csv",

        mime=
        "text/csv"

    )



    word_file=create_word(
        result,
        question
    )


    st.sidebar.download_button(

        label="📄 Prenesi Word",

        data=
        word_file,

        file_name=
        "QAS99_evalvacija.docx",

        mime=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )



    pdf_file=create_pdf(
        result,
        question
    )


    st.sidebar.download_button(

        label="📕 Prenesi PDF",

        data=
        pdf_file,

        file_name=
        "QAS99_evalvacija.pdf",

        mime=
        "application/pdf"

    )
